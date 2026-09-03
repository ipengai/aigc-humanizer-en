import os
import tempfile
import unittest

from docx import Document

from app.helpers.file_output import generate_docx
from app.helpers.segmenter import segment
from app.humanizer.adapter import HumanizerAdapter
from app.text_extract import extract_text, paragraph_list_to_text


try:
    import pymupdf as fitz
except ImportError:  # pragma: no cover - supported dependency fallback
    try:
        import fitz
    except ImportError:
        fitz = None


class _StructuredStub(HumanizerAdapter):
    def humanize(self, text, mode='low', paragraphs=None):
        return self.humanize_structured(text, mode, paragraphs)[0]

    def humanize_structured(self, text, mode='low', paragraphs=None):
        return self._humanize_segmented_structured(
            mode, paragraphs, lambda value: f'Rewritten: {value}'
        )


@unittest.skipIf(fitz is None, 'PyMuPDF is required')
class PdfStructureTests(unittest.TestCase):
    def setUp(self):
        self.tempdir = tempfile.TemporaryDirectory()
        self.pdf_path = os.path.join(self.tempdir.name, 'structured.pdf')
        self._build_pdf(self.pdf_path)

    def tearDown(self):
        self.tempdir.cleanup()

    @staticmethod
    def _build_pdf(path):
        doc = fitz.open()
        cover = doc.new_page(width=595, height=842)
        cover.insert_text((72, 100), 'Research Report', fontsize=22)
        cover.insert_text((72, 160), 'Contents', fontsize=17)
        cover.insert_text((72, 200), '1 Introduction ........ 1', fontsize=11)

        page = doc.new_page(width=595, height=842)
        page.insert_text((72, 90), '1 Introduction', fontsize=17)
        page.insert_textbox(
            fitz.Rect(72, 115, 520, 190),
            'This is the first body paragraph. It contains enough ordinary prose '
            'to establish the dominant body font and remain eligible for rewriting.',
            fontsize=11,
        )

        table_values = [
            ['Metric', 'Value', 'Status'],
            ['Accuracy', '98', 'Good'],
            ['Latency', '120', 'Stable'],
        ]
        for row_index, values in enumerate(table_values):
            y = (250, 274, 310)[row_index]
            for x, value in zip((72, 220, 360), values):
                page.insert_text((x, y), value, fontsize=10)
        # Borderless PDF tables commonly wrap only a later column. Keep the
        # continuation aligned under its cell so extraction must attach it to
        # the Accuracy row rather than emit it as a stray paragraph.
        page.insert_text((360, 286), 'and reliable', fontsize=10)

        pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 40, 30), False)
        pixmap.clear_with(0x4A7EBB)
        page.insert_image(
            fitz.Rect(72, 350, 300, 500), stream=pixmap.tobytes('png')
        )
        page.insert_text((72, 525), 'Figure 1: Example result', fontsize=10)
        page.insert_text((72, 575), 'References', fontsize=17)
        page.insert_text(
            (72, 610), 'Smith, J. (2025). Structured PDF extraction.', fontsize=11
        )
        page.insert_text(
            (72, 635), 'Jones, A. (2024). Document layout analysis.', fontsize=11
        )
        doc.set_toc([
            [1, 'Introduction', 2],
            [1, 'References', 2],
        ])
        doc.save(path)
        doc.close()

    def test_extracts_word_like_structure(self):
        nodes = extract_text(self.pdf_path)

        introduction = next(
            node for node in nodes if node.get('text') == '1 Introduction'
        )
        self.assertTrue(introduction['is_heading'])
        self.assertEqual(introduction['heading_level'], 1)
        self.assertEqual(introduction['style'], 'Heading 1')

        table = next(node for node in nodes if node.get('node_type') == 'table')
        self.assertEqual(table['headers'], ['Metric', 'Value', 'Status'])
        self.assertEqual(
            table['rows'][0], ['Accuracy', '98', 'Good and reliable']
        )

        image = next(node for node in nodes if node.get('node_type') == 'image')
        self.assertEqual(image['page_number'], 2)
        self.assertGreaterEqual(image['display_width'], 200)

        references = [node for node in nodes if node.get('is_reference')]
        self.assertGreaterEqual(len(references), 1)
        self.assertIn('Smith, J.', references[0]['text'])
        self.assertIn('Jones, A.', references[-1]['text'])
        self.assertTrue(all(node['node_type'] == 'reference' for node in references))

        cover = next(node for node in nodes if node.get('text') == 'Research Report')
        self.assertTrue(cover['is_front_matter'])
        self.assertTrue(next(
            node for node in nodes if node.get('text') == 'Contents'
        )['is_toc'])

    def test_nodes_are_preserved_through_segmentation(self):
        nodes = extract_text(self.pdf_path)
        tasks = segment(nodes, mode='median', min_chars=0)
        layout_nodes = [
            task['paragraphs'][0] for task in tasks if task['type'] == 'layout'
        ]
        self.assertTrue(any(node.get('node_type') == 'table' for node in layout_nodes))
        self.assertTrue(any(node.get('node_type') == 'image' for node in layout_nodes))

        _, rewritten_nodes = _StructuredStub().humanize_structured(
            paragraph_list_to_text(nodes), mode='low', paragraphs=nodes
        )
        self.assertTrue(any(
            node.get('node_type') == 'table' for node in rewritten_nodes
        ))
        self.assertTrue(any(
            node.get('node_type') == 'image' for node in rewritten_nodes
        ))
        self.assertTrue(any(
            node.get('node_type') == 'reference' and node.get('is_reference')
            for node in rewritten_nodes
        ))

    def test_docx_export_renders_headings_table_and_image(self):
        nodes = extract_text(self.pdf_path)
        output = generate_docx(
            paragraph_list_to_text(nodes),
            paragraphs=nodes,
            source_path=self.pdf_path,
        )
        document = Document(output)

        self.assertGreaterEqual(len(document.tables), 1)
        self.assertGreaterEqual(len(document.inline_shapes), 1)
        self.assertTrue(any(
            paragraph.style.name == 'Heading 1' and
            paragraph.text == '1 Introduction'
            for paragraph in document.paragraphs
        ))


if __name__ == '__main__':
    unittest.main()
