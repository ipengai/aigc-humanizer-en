"""File output helpers: generate downloadable docx / md / txt responses."""

import io
import logging
import os
from flask import send_file


logger = logging.getLogger(__name__)


def _heading_style_name(level):
    """把标题级别映射为 Word 内置样式名。level 0 为 Title。"""
    if level is None:
        return None
    if level == 0:
        return 'Title'
    return f'Heading {level}'


def _apply_paragraph_style(paragraph, para):
    """根据段落结构标记给段落应用样式（标题/正文/列表/缩进）。"""
    text = para.get('text', '')
    if not text:
        return

    level = para.get('heading_level')

    # 标题样式（Heading 1~9 / Title）
    style_name = _heading_style_name(level)
    if style_name is not None:
        paragraph.style = style_name
        run = paragraph.add_run(text)
        run.bold = bool(para.get('is_bold')) or None
        run.italic = bool(para.get('is_italic')) or None
        return

    marker = para.get('list_text')
    if marker and not text.lstrip().startswith(str(marker)):
        text = f'{marker} {text}'
    run = paragraph.add_run(text)
    run.bold = bool(para.get('is_bold')) or None
    run.italic = bool(para.get('is_italic')) or None

    if para.get('is_caption'):
        try:
            paragraph.style = 'Caption'
        except KeyError:
            pass

    if para.get('is_reference'):
        from docx.shared import Inches
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.first_line_indent = Inches(-0.3)

    # 列表编号：原始编号已还原为文本，直接作为普通文本（保留缩进）
    indent = para.get('indent')
    if indent is not None:
        try:
            # Word 缩进单位 twips（1/20 pt），转 EMU：1 twip = 635 EMU
            paragraph.paragraph_format.left_indent = int(indent) * 635
        except (ValueError, TypeError):
            paragraph.paragraph_format.left_indent = None


def _add_table(doc, node):
    headers = [str(value or '') for value in (node.get('headers') or [])]
    rows = [
        [str(value or '') for value in row]
        for row in (node.get('rows') or [])
    ]
    column_count = max(
        [len(headers)] + [len(row) for row in rows] + [int(node.get('column_count') or 0)]
    )
    if column_count < 1:
        return
    if len(headers) < column_count:
        headers.extend([''] * (column_count - len(headers)))
    normalized_rows = [row + [''] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=1, cols=column_count)
    table.style = 'Table Grid'
    table.autofit = True
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for values in normalized_rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value


def _open_pdf_source(source_path):
    if not source_path or os.path.splitext(source_path)[1].lower() != '.pdf':
        return None
    try:
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz
        return fitz.open(source_path)
    except Exception:
        logger.exception('Failed to open source PDF for image export: %s', source_path)
        return None


def _add_pdf_image(doc, pdf_document, node):
    if pdf_document is None:
        logger.warning(
            'Skipping PDF image because the source copy is unavailable (page=%s)',
            node.get('page_number'),
        )
        return
    page_index = int(node.get('page_number') or 1) - 1
    bbox = node.get('bbox')
    if not bbox or not 0 <= page_index < len(pdf_document):
        return
    try:
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches

        page = pdf_document[page_index]
        clip = fitz.Rect(*bbox) & page.rect
        if clip.is_empty:
            return
        pixmap = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), clip=clip, alpha=False)
        image_stream = io.BytesIO(pixmap.tobytes('png'))
        usable_width = (
            doc.sections[-1].page_width -
            doc.sections[-1].left_margin -
            doc.sections[-1].right_margin
        )
        display_inches = max(0.5, min(float(node.get('display_width') or clip.width) / 72, 6.5))
        width = min(Inches(display_inches), usable_width)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(image_stream, width=width)
    except Exception:
        logger.exception(
            'Failed to render PDF image page=%s bbox=%s',
            node.get('page_number'), bbox,
        )


def generate_docx(text, paragraphs=None, source_path=None):
    """Generate a .docx file in-memory from text content.

    paragraphs: 可选，结构化段落列表 [{'text','is_heading','heading_level','style',...}]。
        提供时按标题级别重建格式（Title / Heading 1~9 / 正文），否则全部按普通段落。
    """
    from docx import Document
    doc = Document()
    pdf_document = _open_pdf_source(source_path)

    try:
        if paragraphs:
            for para in paragraphs:
                node_type = para.get('node_type')
                if node_type == 'table' or ('table' in para and not para.get('text')):
                    _add_table(doc, para)
                    continue
                if node_type == 'image' or ('image' in para and not para.get('text')):
                    _add_pdf_image(doc, pdf_document, para)
                    continue
                ptext = (para.get('text') or '').strip()
                if not ptext:
                    continue
                p = doc.add_paragraph()
                _apply_paragraph_style(p, para)
        else:
            for paragraph in text.split('\n\n'):
                p = doc.add_paragraph(paragraph.strip())
                if not paragraph.strip():
                    p.add_run('\u200b')
    finally:
        if pdf_document is not None:
            pdf_document.close()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_file_response(text, original_format, filename, paragraphs=None,
                           source_path=None):
    """Generate a file response for download based on format."""
    base_name = os.path.splitext(filename)[0] if filename else 'humanized'

    # PDF originals have no layout/fonts preserved, so default to docx output
    if original_format == 'pdf':
        original_format = 'docx'

    if original_format == 'docx':
        buf = generate_docx(
            text, paragraphs=paragraphs, source_path=source_path
        )
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{base_name}_humanized.docx'
        )
    elif original_format == 'md':
        buf = io.BytesIO(text.encode('utf-8'))
        return send_file(
            buf,
            mimetype='text/markdown',
            as_attachment=True,
            download_name=f'{base_name}_humanized.md'
        )
    else:  # txt (default)
        buf = io.BytesIO(text.encode('utf-8'))
        return send_file(
            buf,
            mimetype='text/plain',
            as_attachment=True,
            download_name=f'{base_name}_humanized.txt'
        )
